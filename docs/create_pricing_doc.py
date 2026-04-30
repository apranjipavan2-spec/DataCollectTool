"""
Generate FieldGovern Pricing Strategy Word Document
Run: python create_pricing_doc.py
Output: FieldGovern_Pricing_Strategy.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1A, 0x3C, 0x5E)
MID_BLUE   = RGBColor(0x2E, 0x6D, 0xA4)
ACCENT     = RGBColor(0x00, 0xA8, 0xE8)
LIGHT_GREY = RGBColor(0xF2, 0xF4, 0xF7)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x1A, 0x1A, 0x1A)
KOBO_RED   = RGBColor(0xC0, 0x39, 0x2B)
FG_GREEN   = RGBColor(0x27, 0xAE, 0x60)
AMBER      = RGBColor(0xE6, 0x7E, 0x22)

def set_cell_bg(cell, hex_color: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{hex_color[0]:02X}{hex_color[1]:02X}{hex_color[2]:02X}')
    tcPr.append(shd)

def rgb_hex(c: RGBColor):
    return f'{c[0]:02X}{c[1]:02X}{c[2]:02X}'

def h1(text):
    p  = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_BLUE
    return p

def h2(text):
    p  = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = MID_BLUE
    return p

def h3(text):
    p  = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    return p

def body(text, italic=False, color=None):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size   = Pt(10)
    run.italic      = italic
    if color:
        run.font.color.rgb = color
    return p

def note(text):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(f'⚠  {text}')
    run.font.size      = Pt(9)
    run.font.color.rgb = AMBER
    run.italic         = True
    return p

def divider():
    p  = doc.add_paragraph('─' * 90)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.runs[0]
    run.font.size      = Pt(7)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

def make_table(headers, rows, col_widths=None, header_bg=None):
    """Create a styled table."""
    n_cols = len(headers)
    tbl    = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = 'Table Grid'

    hdr_bg = header_bg or DARK_BLUE

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_bg(cell, hdr_bg)
        p   = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold           = True
        run.font.size      = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment        = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        alt   = (r_idx % 2 == 1)
        for c_idx, val in enumerate(row):
            cell = cells[c_idx]
            if alt:
                set_cell_bg(cell, RGBColor(0xF7, 0xF9, 0xFC))
            p   = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if c_idx == 0:
                run.bold = True

    # Column widths
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return tbl

# ═══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('FieldGovern')
run.bold           = True
run.font.size      = Pt(32)
run.font.color.rgb = DARK_BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Pricing Strategy Document')
run2.bold           = True
run2.font.size      = Pt(18)
run2.font.color.rgb = MID_BLUE

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Segmented Pricing · Competitive Benchmarking · Feature Matrix')
run3.font.size      = Pt(11)
run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run3.italic         = True

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run('DRAFT — For Internal Review & Edits')
run4.bold           = True
run4.font.size      = Pt(10)
run4.font.color.rgb = AMBER

doc.add_paragraph()
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5 = p5.add_run('Version 1.0  |  April 2026  |  Confidential')
run5.font.size      = Pt(9)
run5.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — PRICING PHILOSOPHY
# ═══════════════════════════════════════════════════════════════════════════════
h1('1. Pricing Philosophy')
body(
    'FieldGovern is positioned as the AI-native alternative to KoboToolbox — not as a cheap substitute. '
    'Our pricing is set at 80–90% of KoboToolbox\'s equivalent plan to signal quality and confidence while '
    'offering materially superior value through built-in AI capabilities that KoboToolbox does not provide at any tier.'
)
body(
    'Key principle: Every FieldGovern paid plan includes features that KoboToolbox users currently pay '
    'for separately (data analyst salary, PowerBI licences, report writing time). Our pricing reflects the '
    'total-cost-of-ownership advantage, not just the licence fee.'
)

h2('1.1  KoboToolbox Reference Pricing (Benchmark)')
note('Prices sourced from kobotoolbox.org/pricing — last verified February 2025. Verify before finalising FieldGovern pricing.')

kobo_rows = [
    ['Starter (Standard)',          '$21/mo',   '1,000/mo',    '0.5 GB',    'Unlimited', '5 min/mo ASR · 3K chars translation · Community support only'],
    ['Community (Standard)',        '$88/mo',   '5,000/mo',    '1 GB',      'Unlimited', '10 min/mo ASR · 6K chars translation · Community support only'],
    ['Professional (Standard)',     '$166/mo',  '25,000/mo\nor 300K/yr', 'Unlimited', 'Unlimited', '2FA · 120 min/mo ASR · 72K chars translation · Direct support'],
    ['Enterprise (Standard)',       'Custom',   'Custom',      'Unlimited', 'Unlimited', 'SSO · Dedicated support · Expert consultations'],
    ['Community Nonprofit',         'FREE',     '5,000/mo',    '1 GB',      'Unlimited', '10 min/mo ASR · 6K chars · Community support only'],
    ['Professional (Nonprofit)',    '$129/mo',  '25,000/mo\nor 300K/yr', 'Unlimited', 'Unlimited', '2FA · 120 min/mo ASR · 72K chars translation · Direct support'],
    ['Enterprise (Nonprofit)',      'Custom',   'Custom',      'Unlimited', 'Unlimited', 'SSO · Dedicated support · Expert consultations · Roadmap input'],
]
make_table(
    ['Plan', 'Price', 'Submissions', 'Storage', 'Forms/Projects', 'Key Features'],
    kobo_rows,
    col_widths=[3.5, 2.0, 2.5, 2.0, 2.5, 6.0],
    header_bg=KOBO_RED
)

h2('1.2  Free Tier Baseline — All Segments')
body(
    'Every FieldGovern segment includes a Free tier. This matches KoboToolbox\'s Starter limits '
    '(1,000 submissions, 500 MB storage) but adds superior UX and includes one AI-generated report per month '
    'as a conversion hook.'
)
body('Free tier features (applies to ALL segments):')

free_features = [
    '1,000 survey submissions per month',
    '500 MB file storage',
    'Unlimited forms and projects',
    'Unlimited data collectors and collaborators',
    'Online and offline data collection (mobile + web)',
    'All 25+ question types (text, select, GPS, photo, audio, video, barcode, signature, matrix, ranking, calculation)',
    'Advanced skip logic and branching',
    'Cascading select questions',
    'Nested repeat groups',
    'Form validation and calculation rules',
    'Multi-language form support',
    'XLSForm import and export',
    'GPS capture and map view',
    'Media capture (photo, audio, video)',
    'Data table view, filtering, and search',
    'Basic reports and charts (bar, pie, line)',
    'Data export: CSV, Excel, GeoJSON, KML',
    'REST API access (read-only)',
    'Automatic speech-to-text: 5 minutes per month',
    'Machine translation: 3,000 characters per month',
    '1 AI-generated summary report per month (FieldGovern exclusive)',
    'Community forum and self-serve documentation',
]
for f in free_features:
    p = doc.add_paragraph(f, style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.runs[0].font.size = Pt(9)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — SEGMENT PRICING
# ═══════════════════════════════════════════════════════════════════════════════
h1('2. Segmented Pricing Plans')
note('All INR prices shown at ₹83 per USD. Round to nearest psychologically attractive price point before publishing.')
note('Annual billing: 2 months free (≈17% discount). Recommended default billing cycle for all segments.')

# ─────────────────────────────────────────────────────────────────────────────
# SEGMENT A — NGO
# ─────────────────────────────────────────────────────────────────────────────
doc.add_paragraph()
h2('Segment A — NGO / Social Impact / M&E')
body(
    'Primary buyers: M&E officers, program directors, data managers. '
    'Budget source: grant funding (typically 6–12 month cycles). '
    'Core pain: Kobo collects data but they still pay ₹30–50K/month for a data analyst + report writer. '
    'Our pitch: FieldGovern replaces those hidden costs at a fraction of the total.'
)

note('Kobo comparison: Free Nonprofit (₹0, 5K sub) · Professional Nonprofit ($129/mo). '
     'FieldGovern targets the paid Kobo NGO segment at 80–90% of cost with AI built in.')

ngo_plans = [
    ['Free', '₹0\n($0)', '1,000/mo', '500 MB', '—', '—',
     '• All standard free features (see §1.2)\n• 1 AI summary report/month\n• Community support'],

    ['Starter', '₹6,499/mo\n(~$78)\n[Kobo: $88 Community]', '5,000/mo', '1 GB', '—', 'Included',
     '• All Free features\n• AI Data Cleaning with audit trail\n• Unlimited FG Writer reports\n'
     '• Auto speech-to-text: 10 min/month\n• Machine translation: 6,000 chars/month\n'
     '• Email support (48hr SLA)\n• Form logic templates for M&E\n• XLSForm bulk import'],

    ['Growth', '₹10,999/mo\n(~$132)\n[Kobo: $129 Pro Nonprofit]', '25,000/mo\nor 3,00,000/yr',
     'Unlimited', '2FA included', 'Included',
     '• All Starter features\n• Smart Builder (AI table design)\n• AI Interpretation per table\n'
     '• Iterative Refinement (AI re-generation)\n• Panel Study & Attrition Reports\n'
     '• Auto speech-to-text: 120 min/month\n• Machine translation: 72,000 chars/month\n'
     '• Two-factor Authentication (2FA)\n• REST API (read + write)\n• Webhooks/REST services\n'
     '• Direct user support (24hr SLA)\n• Donor report templates (USAID, EU, Gates)'],

    ['Enterprise', 'Custom\n[Kobo: Custom]', 'Custom', 'Unlimited', 'SSO + 2FA', 'Custom',
     '• All Growth features\n• Multi-org / consortium dashboard\n• Single Sign-On (SSO)\n'
     '• Custom submission volume\n• Custom transcription/translation quotas\n'
     '• On-premise deployment (DPDP compliance)\n• Dedicated Customer Success Manager\n'
     '• Custom onboarding and training\n• Priority SLA (4hr response)\n'
     '• White-label report output\n• Audit log access\n• Roadmap input'],
]
make_table(
    ['Plan', 'Price', 'Submissions', 'Storage', 'Security', 'AI Cleaning & Writer', 'Features Included'],
    ngo_plans,
    col_widths=[2.0, 3.0, 2.5, 2.0, 2.5, 2.5, 7.5],
    header_bg=FG_GREEN
)

note('REVIEWER NOTE: Should we offer a 6-month billing cycle for NGOs aligned to grant tranches? '
     'Kobo does not offer this — it could be a meaningful differentiator.')

# ─────────────────────────────────────────────────────────────────────────────
# SEGMENT B — GOVERNMENT
# ─────────────────────────────────────────────────────────────────────────────
doc.add_paragraph()
h2('Segment B — Government (State + Central Schemes)')
body(
    'Primary buyers: IT departments, programme implementation units (PIUs), district collectors\' offices. '
    'Budget source: scheme budgets (annual). Procurement: GEM portal (GeM), L1 tendering. '
    'Core pain: Custom-built MIS systems cost ₹30–80 lakh upfront + ₹5–10L/year maintenance. '
    'Our pitch: deploy in days, not 18 months; DPDP-compliant; AI-verified data quality built in.'
)

note('Kobo comparison: Government orgs typically fall on Kobo Professional ($166/mo) or self-host. '
     'FieldGovern\'s real competitor here is custom development — not Kobo.')

gov_plans = [
    ['Free', '₹0\n($0)', '1,000/mo', '500 MB', '—',
     '• All standard free features (see §1.2)\n• 1 AI summary report/month'],

    ['Department', '₹13,799/mo\n(~$166)\n[Kobo Professional: $166]', '25,000/mo\nor 3,00,000/yr',
     'Unlimited', '2FA',
     '• All Free features\n• AI Data Cleaning with full audit trail\n• Unlimited FG Writer reports\n'
     '• GPS-based fraud detection flags\n• Role-based access control (RBAC)\n'
     '• Audit log (all data changes timestamped)\n• Auto speech-to-text: 120 min/month\n'
     '• Machine translation: 72,000 chars/month\n• Two-factor Authentication (2FA)\n'
     '• REST API + Webhooks\n• Direct user support (24hr SLA)\n'
     '• Standard scheme report templates'],

    ['State / Multi-Scheme', '₹24,999/mo\n(~$301)', '1,00,000/mo', 'Unlimited', '2FA + Roles',
     '• All Department features\n• Multi-scheme unified dashboard\n• Smart Builder + AI Interpret\n'
     '• Panel Study (multi-wave beneficiary tracking)\n• Advanced RBAC (district/block/GP hierarchy)\n'
     '• Real-time field agent monitoring\n• Offline-first mobile app (PWA)\n'
     '• GPS spoofing detection\n• Priority support (8hr SLA)\n'
     '• Vernacular language UI (Hindi roadmap)'],

    ['National / Enterprise', 'Custom\n(₹20L+/year typical)', 'Custom', 'Unlimited', 'SSO + 2FA + Audit',
     '• All State features\n• Single Sign-On (SSO)\n• On-premise / private cloud deployment\n'
     '• DPDP Act compliance framework\n• NIC/NICNET integration support\n'
     '• Custom data retention and purge policies\n• GeM-compatible contract structure\n'
     '• White-label (ministry/department branding)\n• Dedicated success team\n'
     '• Custom training for field agents\n• Priority SLA (2hr response)\n'
     '• National-level aggregation dashboard'],
]
make_table(
    ['Plan', 'Price', 'Submissions', 'Storage', 'Security', 'Features Included'],
    gov_plans,
    col_widths=[2.5, 3.2, 2.5, 2.0, 2.5, 8.3],
    header_bg=DARK_BLUE
)

note('REVIEWER NOTE: GeM listing is mandatory to sell to government. Plan for GeM onboarding before targeting this segment. '
     'Annual-only billing recommended for government (matches procurement cycles).')

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SEGMENT C — RESEARCH
# ─────────────────────────────────────────────────────────────────────────────
h2('Segment C — Research Organizations (Universities, Think Tanks, ICMR, ICAR)')
body(
    'Primary buyers: Principal investigators, M&E leads, data managers. '
    'Budget source: research grants (ICMR, DST, USAID, Ford, Gates), institutional budgets. '
    'Core pain: SurveyCTO ($299/month) is expensive and has no AI; Kobo has no panel study depth or '
    'statistical export. Our pitch: publication-grade data quality + AI cleaning + panel study + SPSS/Stata export '
    'at 80% of SurveyCTO price.'
)

note('Kobo comparison: Research orgs pay $129 (Professional Nonprofit) or $166 (Professional). '
     'SurveyCTO is the premium benchmark at $299/month.')

res_plans = [
    ['Free / Academic', '₹0\n($0)', '1,000/mo', '500 MB', '—',
     '• All standard free features (see §1.2)\n• Suitable for student dissertations and pilot surveys\n• 1 AI summary report/month'],

    ['Research', '₹13,799/mo\n(~$166)\n[Kobo Pro: $166 | SurveyCTO: $299]',
     '25,000/mo\nor 3,00,000/yr', 'Unlimited', '2FA',
     '• All Free features\n• AI Data Cleaning with full cleaning audit trail\n'
     '• Unlimited FG Writer reports\n• Panel Study with attrition tracking\n'
     '• Multi-wave longitudinal data management\n• SPSS and Stata-compatible export\n'
     '• Data dictionary export (codebook)\n• GeoJSON and KML export\n'
     '• Auto speech-to-text: 120 min/month\n• Machine translation: 72,000 chars/month\n'
     '• Two-factor Authentication (2FA)\n• REST API + Webhooks\n'
     '• Direct user support (24hr SLA)'],

    ['Institutional', '₹20,799/mo\n(~$250)', '75,000/mo', 'Unlimited', '2FA + SSO option',
     '• All Research features\n• Smart Builder + AI Interpretation\n'
     '• Multi-PI sub-accounts (research assistants with role limits)\n'
     '• IRB/ethics audit trail (timestamped access log)\n'
     '• Advanced panel study (nested waves, attrition cohorts)\n'
     '• Custom data retention policies\n• Priority support (12hr SLA)\n'
     '• SSO option (university Microsoft/Google SSO)'],

    ['University Enterprise', 'Custom', 'Unlimited', 'Unlimited', 'SSO + 2FA + Audit',
     '• All Institutional features\n• University-wide licence (all departments)\n'
     '• On-premise for sensitive health/social data\n• Dedicated onboarding for research teams\n'
     '• Custom data anonymisation pipeline\n• Priority SLA (4hr response)\n'
     '• Roadmap input on research-specific features'],
]
make_table(
    ['Plan', 'Price', 'Submissions', 'Storage', 'Security', 'Features Included'],
    res_plans,
    col_widths=[2.5, 3.5, 2.5, 2.0, 2.5, 8.0],
    header_bg=MID_BLUE
)

note('REVIEWER NOTE: Should academic/student accounts get a separate verified free tier '
     '(e.g., .edu/.ac.in email verification, 2,000 submissions instead of 1,000)? '
     'This could drive university-wide adoption faster.')

# ─────────────────────────────────────────────────────────────────────────────
# SEGMENT D — CORPORATE
# ─────────────────────────────────────────────────────────────────────────────
doc.add_paragraph()
h2('Segment D — Corporate / Private Sector (FMCG, MFI/NBFC, Agri-Tech, ESG/Climate)')
body(
    'Primary buyers: Research heads, field operations managers, compliance officers. '
    'Budget source: operational budget (OpEx). Highest willingness to pay. '
    'Core pain: Qualtrics ($1,500+/month) is overkill for field collection; Kobo has no AI or real-time dashboard. '
    'Our pitch: field-grade collection + AI analysis + photo/GPS audit at 10% of Qualtrics cost.'
)

note('Kobo comparison: Corporate users pay $166 (Professional) or Enterprise custom. '
     'Qualtrics is the enterprise benchmark at $1,500+/month.')

corp_plans = [
    ['Free', '₹0\n($0)', '1,000/mo', '500 MB', '—',
     '• All standard free features (see §1.2)\n• 1 AI summary report/month\n• Suitable for pilots and PoCs'],

    ['Business', '₹13,799/mo\n(~$166)\n[Kobo Professional: $166]',
     '25,000/mo\nor 3,00,000/yr', 'Unlimited', '2FA',
     '• All Free features\n• AI Data Cleaning with audit trail\n• Unlimited FG Writer reports\n'
     '• Smart Builder (AI-designed cross-tabs)\n• GPS fraud detection and spoofing flags\n'
     '• Photo audit support (field agent photo verification)\n'
     '• REST API (read + write) + Webhooks\n• Auto speech-to-text: 120 min/month\n'
     '• Machine translation: 72,000 chars/month\n• Two-factor Authentication (2FA)\n'
     '• Direct user support (24hr SLA)\n• Excel export with auto-formatting'],

    ['Corporate', '₹20,799/mo\n(~$250)', '2,00,000/mo', 'Unlimited', '2FA + Roles + Audit',
     '• All Business features\n• AI Interpretation per tabulation\n'
     '• Multi-branch / multi-district unified dashboard\n• Real-time field agent monitoring\n'
     '• Teams management (unlimited members)\n• Advanced RBAC (zone/region/branch hierarchy)\n'
     '• Offline-first mobile app (PWA)\n• Priority support (8hr SLA)\n'
     '• ESG metric calculation (for climate/sustainability teams)\n'
     '• Custom report branding (logo, colours)'],

    ['Enterprise', 'Custom', 'Custom', 'Unlimited', 'SSO + 2FA + Full Audit',
     '• All Corporate features\n• Single Sign-On (SSO)\n• White-label platform\n'
     '• On-premise / private cloud deployment\n• CRM/ERP integration support (REST API)\n'
     '• Dedicated Customer Success Manager\n• Custom onboarding and training\n'
     '• Priority SLA (2hr response)\n• Custom AI model fine-tuning (roadmap)\n'
     '• Advanced fraud detection dashboard (MFI/NBFC use case)'],
]
make_table(
    ['Plan', 'Price', 'Submissions', 'Storage', 'Security', 'Features Included'],
    corp_plans,
    col_widths=[2.5, 3.2, 2.5, 2.0, 2.8, 8.0],
    header_bg=AMBER
)

note('REVIEWER NOTE: MFI/NBFC segment may need a specific KYC + GPS spoofing detection bundle. '
     'Consider a vertical add-on pricing model (₹3,000/month on top of Business plan) rather than a full separate tier.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — FEATURE COMPARISON MATRIX
# ═══════════════════════════════════════════════════════════════════════════════
h1('3. FieldGovern vs KoboToolbox — Full Feature Comparison')
body('✓ = Included on all plans  |  P = Paid plans only  |  E = Enterprise only  |  ✕ = Not available  |  ★ = FieldGovern exclusive')

TICK  = '✓'
CROSS = '✕'
STAR  = '★'
PAY   = 'Paid'
ENT   = 'Enterprise'

comparison = [
    # Category | Feature | Kobo Free | Kobo Paid | Kobo Enterprise | FG Free | FG Paid | FG Enterprise
    ['FORM BUILDING',       'Drag-and-drop form builder',                TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    '25+ question types',                         TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Skip logic / branching',                     TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Cascading select questions',                  TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Nested repeat groups',                        TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Calculation and validation rules',            TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'XLSForm import / export',                     TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Multi-language form support',                 TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Guidance hints and notes',                    TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Form versioning',                             TICK, TICK, TICK,   TICK, TICK, TICK],
    ['DATA COLLECTION',     'Online (web browser)',                        TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Offline (mobile app / PWA)',                  TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'GPS / location capture',                      TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Photo capture',                               TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Audio / video capture',                       TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Barcode / QR code scanning',                  TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Digital signature capture',                   TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Biometric / face capture',                    CROSS,CROSS,CROSS,  CROSS,CROSS,'Roadmap'],
    ['DATA MANAGEMENT',     'Data table view and filtering',               TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Edit submitted records',                      TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Validation / moderation queue',               TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Photo / media gallery',                       TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Map view for GPS data',                       TICK, TICK, TICK,   TICK, TICK, TICK],
    ['EXPORT & REPORTING',  'CSV export',                                  TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Excel (XLSX) export',                         TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'GeoJSON / KML export',                        TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'SPSS / Stata export',                         CROSS,CROSS,CROSS,  CROSS,PAY,  TICK],
    ['',                    'Custom report charts (bar, pie, line)',        TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Cross-tabulation (basic)',                     TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Data dictionary / codebook export',            CROSS,CROSS,CROSS,  CROSS,PAY,  TICK],
    ['AI FEATURES',         'AI Data Cleaning',                            CROSS,CROSS,CROSS,  CROSS,STAR, STAR],
    ['',                    'AI Cleaning audit trail',                     CROSS,CROSS,CROSS,  CROSS,STAR, STAR],
    ['',                    'FG Writer (AI report generation)',             CROSS,CROSS,CROSS,  '1/mo',STAR,STAR],
    ['',                    'Smart Builder (AI table design)',              CROSS,CROSS,CROSS,  CROSS,STAR, STAR],
    ['',                    'AI Interpretation per tabulation',            CROSS,CROSS,CROSS,  CROSS,STAR, STAR],
    ['',                    'Iterative Refinement (AI re-generation)',      CROSS,CROSS,CROSS,  CROSS,STAR, STAR],
    ['',                    'Donor report templates (USAID, EU, etc.)',    CROSS,CROSS,CROSS,  CROSS,STAR, STAR],
    ['',                    'Automatic speech-to-text (ASR)',              '5 min','120 min','Custom','5 min','120 min','Custom'],
    ['',                    'Machine translation',                          '3K chars','72K chars','Custom','3K chars','72K chars','Custom'],
    ['INTEGRATION',         'REST API access',                             TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'REST API (write access)',                      TICK, TICK, TICK,   CROSS,TICK, TICK],
    ['',                    'Webhooks / REST services',                    TICK, TICK, TICK,   CROSS,TICK, TICK],
    ['',                    'Power BI / Tableau connector',                CROSS,CROSS,CROSS,  CROSS,CROSS,'Roadmap'],
    ['',                    'Zapier / n8n integration',                    TICK, TICK, TICK,   CROSS,TICK, TICK],
    ['',                    'CRM / ERP integration',                       CROSS,CROSS,ENT,    CROSS,CROSS,TICK],
    ['SECURITY & ADMIN',    'Two-factor authentication (2FA)',              CROSS,TICK, TICK,   CROSS,TICK, TICK],
    ['',                    'Single Sign-On (SSO)',                        CROSS,CROSS,TICK,   CROSS,CROSS,TICK],
    ['',                    'Audit log (data change history)',              CROSS,CROSS,CROSS,  CROSS,PAY,  TICK],
    ['',                    'GPS spoofing / fraud detection',               CROSS,CROSS,CROSS,  CROSS,STAR, STAR],
    ['',                    'Role-based access control (RBAC)',            TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Advanced RBAC (hierarchy: region/district)',  CROSS,CROSS,TICK,   CROSS,TICK, TICK],
    ['',                    'On-premise / private cloud deployment',       CROSS,CROSS,TICK,   CROSS,CROSS,TICK],
    ['',                    'DPDP Act compliance framework',               CROSS,CROSS,CROSS,  CROSS,CROSS,STAR],
    ['COLLABORATION',       'Unlimited data collectors',                   TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Teams / group management',                    CROSS,PAY,  TICK,   CROSS,TICK, TICK],
    ['',                    'Multi-org / consortium dashboard',            CROSS,CROSS,TICK,   CROSS,CROSS,TICK],
    ['',                    'Panel study (longitudinal)',                  CROSS,CROSS,CROSS,  CROSS,STAR, STAR],
    ['',                    'Attrition / retention reports',               CROSS,CROSS,CROSS,  CROSS,STAR, STAR],
    ['SUPPORT',             'Community forum',                             TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Self-serve documentation',                    TICK, TICK, TICK,   TICK, TICK, TICK],
    ['',                    'Direct user support',                         CROSS,TICK, TICK,   CROSS,TICK, TICK],
    ['',                    'Priority SLA support',                        CROSS,CROSS,TICK,   CROSS,CROSS,TICK],
    ['',                    'Dedicated Customer Success Manager',          CROSS,CROSS,TICK,   CROSS,CROSS,TICK],
    ['',                    'Custom onboarding and training',              CROSS,CROSS,TICK,   CROSS,CROSS,TICK],
]

cmp_tbl = doc.add_table(rows=1 + len(comparison), cols=8)
cmp_tbl.style = 'Table Grid'

# Header
hdrs = ['Category', 'Feature',
        'Kobo Free', 'Kobo Paid', 'Kobo Ent.',
        'FG Free', 'FG Paid', 'FG Ent.']
for i, h in enumerate(hdrs):
    cell = cmp_tbl.rows[0].cells[i]
    set_cell_bg(cell, DARK_BLUE)
    p   = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True; run.font.size = Pt(8); run.font.color.rgb = WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

current_cat = ''
for r_idx, row in enumerate(comparison):
    cells = cmp_tbl.rows[r_idx + 1].cells
    cat, feat = row[0], row[1]
    vals_kobo = row[2:5]
    vals_fg   = row[5:8]

    # Category label
    if cat and cat != current_cat:
        current_cat = cat
        set_cell_bg(cells[0], RGBColor(0xE8, 0xF0, 0xF8))
        run = cells[0].paragraphs[0].add_run(cat)
        run.bold = True; run.font.size = Pt(8); run.font.color.rgb = DARK_BLUE
    else:
        cells[0].paragraphs[0].add_run('')

    # Feature name
    run = cells[1].paragraphs[0].add_run(feat)
    run.font.size = Pt(8)
    if r_idx % 2 == 1:
        set_cell_bg(cells[1], RGBColor(0xF7, 0xF9, 0xFC))

    # Kobo values
    for c_i, val in enumerate(vals_kobo, start=2):
        cell = cells[c_i]
        if r_idx % 2 == 1:
            set_cell_bg(cell, RGBColor(0xFF, 0xF5, 0xF5))
        run = cell.paragraphs[0].add_run(str(val))
        run.font.size = Pt(8)
        if val == CROSS:
            run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
        elif val == TICK:
            run.font.color.rgb = FG_GREEN
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # FG values
    for c_i, val in enumerate(vals_fg, start=5):
        cell = cells[c_i]
        if r_idx % 2 == 1:
            set_cell_bg(cell, RGBColor(0xF0, 0xFB, 0xF5))
        run = cell.paragraphs[0].add_run(str(val))
        run.font.size = Pt(8)
        if val == CROSS:
            run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
        elif val in (TICK, STAR):
            run.font.color.rgb = FG_GREEN
        if STAR in str(val):
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Column widths
widths = [2.0, 5.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5]
for row in cmp_tbl.rows:
    for i, w in enumerate(widths):
        row.cells[i].width = Cm(w)

doc.add_paragraph()
body('★ = FieldGovern exclusive feature not available on KoboToolbox at any plan tier.', italic=True, color=FG_GREEN)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — PRICING SUMMARY TABLE
# ═══════════════════════════════════════════════════════════════════════════════
h1('4. Pricing Summary — All Segments at a Glance')

summary_rows = [
    ['NGO — Free',             '₹0',           '$0',    '1,000/mo',    '500 MB',     'Kobo: Free nonprofit (5K sub)'],
    ['NGO — Starter',          '₹6,499/mo',    '~$78',  '5,000/mo',    '1 GB',       'Kobo Community: $88/mo'],
    ['NGO — Growth',           '₹10,999/mo',   '~$132', '25K/mo',      'Unlimited',  'Kobo Pro Nonprofit: $129/mo'],
    ['NGO — Enterprise',       'Custom',        'Custom','Custom',      'Unlimited',  'Kobo Nonprofit Enterprise: Custom'],
    ['—', '—', '—', '—', '—', '—'],
    ['Govt — Free',            '₹0',           '$0',    '1,000/mo',    '500 MB',     'Same free tier'],
    ['Govt — Department',      '₹13,799/mo',   '~$166', '25K/mo',      'Unlimited',  'Kobo Professional: $166/mo'],
    ['Govt — State',           '₹24,999/mo',   '~$301', '1L/mo',       'Unlimited',  'Custom / no direct comparison'],
    ['Govt — National Ent.',   'Custom',        'Custom','Custom',      'Unlimited',  'Kobo Enterprise: Custom'],
    ['—', '—', '—', '—', '—', '—'],
    ['Research — Free',        '₹0',           '$0',    '1,000/mo',    '500 MB',     'Same free tier'],
    ['Research — Research',    '₹13,799/mo',   '~$166', '25K/mo',      'Unlimited',  'SurveyCTO: $299/mo | Kobo Pro: $166'],
    ['Research — Institutional','₹20,799/mo',  '~$250', '75K/mo',      'Unlimited',  'Between Kobo Pro and Enterprise'],
    ['Research — Univ. Ent.',  'Custom',        'Custom','Custom',      'Unlimited',  'Kobo Enterprise: Custom'],
    ['—', '—', '—', '—', '—', '—'],
    ['Corporate — Free',       '₹0',           '$0',    '1,000/mo',    '500 MB',     'Same free tier'],
    ['Corporate — Business',   '₹13,799/mo',   '~$166', '25K/mo',      'Unlimited',  'Kobo Professional: $166/mo'],
    ['Corporate — Corporate',  '₹20,799/mo',   '~$250', '2L/mo',       'Unlimited',  'Qualtrics: $1,500+/mo'],
    ['Corporate — Enterprise', 'Custom',        'Custom','Custom',      'Unlimited',  'Kobo Enterprise: Custom'],
]
make_table(
    ['Segment & Plan', 'INR Price', 'USD Equiv.', 'Submissions', 'Storage', 'Kobo Benchmark'],
    summary_rows,
    col_widths=[3.8, 2.5, 2.2, 2.5, 2.2, 5.8],
    header_bg=DARK_BLUE
)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — REVIEWER NOTES
# ═══════════════════════════════════════════════════════════════════════════════
h1('5. Open Questions for Reviewer')
body('Please review and suggest changes directly in this document. Key decisions still to be made:')

questions = [
    '1. PRICES: Are the INR price points psychologically right? (e.g., ₹13,799 vs ₹14,999 vs ₹12,999?) '
       'Suggest your preferred price for each plan.',
    '2. NGO 6-MONTH BILLING: Should we offer 6-month billing aligned to grant cycles? What discount? '
       '(Kobo does not offer this — potential differentiator.)',
    '3. ACADEMIC TIER: Should the Research Free plan have a verified academic sub-tier '
       '(.edu/.ac.in email) with 2,000 submissions instead of 1,000?',
    '4. GOVERNMENT ANNUAL-ONLY: Should government plans be annual-only (matching procurement cycles), '
       'or also available monthly at a 20% premium?',
    '5. MFI VERTICAL ADD-ON: Should we create a specific MFI/NBFC add-on bundle '
       '(GPS spoofing + KYC features) at ₹3,000/month on top of Business, '
       'rather than a full separate plan?',
    '6. FREE TIER LIMIT: 1,000 submissions/month — is this too restrictive for NGO pilots? '
       'Kobo Free gives 5,000. Should our Free give 1,000 but with 1 AI report as the hook?',
    '7. AI FEATURES ON FREE: Currently 1 AI report/month on Free tier. '
       'Should we increase to 3/month or keep at 1 to drive upgrade?',
    '8. CUSTOM DOMAIN / WHITE-LABEL: At what tier should white-label reports be available? '
       'Currently Enterprise-only. Should Growth NGO also get white-label for donor reports?',
    '9. STORAGE: Is 500 MB on Free too small for photo-heavy surveys? '
       'Kobo Starter is also 0.5 GB — should we match or differentiate?',
    '10. CURRENCY: Should we publish prices in USD or INR on the pricing page? '
        'Or dual-currency (INR primary, USD secondary)?',
]
for q in questions:
    p = doc.add_paragraph(q, style='List Bullet')
    p.paragraph_format.space_after = Pt(5)
    p.runs[0].font.size = Pt(10)

doc.add_paragraph()
divider()
body('Sources: kobotoolbox.org/pricing · saasworthy.com/product/kobotoolbox/pricing · '
     'g2.com/products/kobotoolbox/pricing', italic=True, color=RGBColor(0x99, 0x99, 0x99))
body('Document prepared: April 2026  |  FieldGovern Internal Strategy',
     italic=True, color=RGBColor(0x99, 0x99, 0x99))

# ── Save ─────────────────────────────────────────────────────────────────────
out = r'C:\Life\DataCollectTool\docs\FieldGovern_Pricing_Strategy.docx'
doc.save(out)
print(f'Saved: {out}')
