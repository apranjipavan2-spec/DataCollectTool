"""Codebook export (Phase 6).

Generates a DOCX data dictionary from the dataset + column_roles metadata.
Each variable gets: name, type, scale, role, value labels, n_missing, n_unique,
descriptive summary (numeric: mean/SD/range; categorical: frequencies).

Output mirrors the standard codebook structure used in survey research so the
file can be attached as a methodology appendix to a report.
"""

from __future__ import annotations

import traceback
from datetime import datetime
from io import BytesIO

import numpy as np
import pandas as pd
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..shared import (
    datasets,
    column_roles,
    study_designs,
    column_type_overrides,
    add_audit_log,
    EXPORTS_DIR,
)

router = APIRouter()


class CodebookConfig(BaseModel):
    dataset_id: str
    filename: Optional[str] = None
    include_frequencies: bool = True
    max_freq_levels: int = 25
    title: Optional[str] = None
    subtitle: Optional[str] = None


def _missing_count(s: pd.Series) -> int:
    return int(s.isna().sum())


def _detect_type(s: pd.Series, override: Optional[str]) -> str:
    if override:
        return override
    if pd.api.types.is_numeric_dtype(s):
        return "numeric"
    if pd.api.types.is_datetime64_any_dtype(s):
        return "date"
    return "text"


def _numeric_summary(s: pd.Series) -> dict:
    s2 = pd.to_numeric(s, errors="coerce").dropna()
    if len(s2) == 0:
        return {"n": 0}
    return {
        "n": int(len(s2)),
        "mean": float(s2.mean()),
        "sd": float(s2.std(ddof=1)) if len(s2) > 1 else 0.0,
        "min": float(s2.min()),
        "max": float(s2.max()),
        "median": float(s2.median()),
        "q1": float(s2.quantile(0.25)),
        "q3": float(s2.quantile(0.75)),
    }


def _frequency(s: pd.Series, max_levels: int = 25) -> list[tuple[str, int, float]]:
    vc = s.dropna().value_counts()
    total = vc.sum() or 1
    out = []
    for k, v in vc.head(max_levels).items():
        out.append((str(k), int(v), float(v) / total * 100.0))
    if len(vc) > max_levels:
        rest_n = int(vc.iloc[max_levels:].sum())
        out.append((f"… {len(vc) - max_levels} more", rest_n, rest_n / total * 100.0))
    return out


@router.post("/api/export/codebook")
async def export_codebook(config: CodebookConfig):
    if config.dataset_id not in datasets:
        raise HTTPException(404, "Dataset not found")

    try:
        ds = datasets[config.dataset_id]
        df: pd.DataFrame = ds.get("df") if isinstance(ds, dict) else ds
        if not isinstance(df, pd.DataFrame):
            df = ds["df"]
        roles = column_roles.get(config.dataset_id, {}) or {}
        overrides = column_type_overrides.get(config.dataset_id, {}) or {}
        design = study_designs.get(config.dataset_id, {}) or {}

        doc = Document()

        # ── Title page ──
        h = doc.add_heading(config.title or "Data Codebook", level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if config.subtitle:
            sub = doc.add_paragraph(config.subtitle)
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub.runs[0].font.size = Pt(12)
            sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # ── Dataset summary ──
        doc.add_heading("Dataset summary", level=1)
        s_table = doc.add_table(rows=0, cols=2)
        s_table.style = "Light Grid Accent 1"
        for key, val in [
            ("Rows", f"{len(df):,}"),
            ("Variables", f"{len(df.columns)}"),
            ("Date generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ]:
            r = s_table.add_row().cells
            r[0].text = key
            r[1].text = str(val)

        # ── Study design block ──
        if design:
            doc.add_heading("Study design", level=1)
            d_table = doc.add_table(rows=0, cols=2)
            d_table.style = "Light Grid Accent 1"
            for key in ("design_type", "treatment_col", "treatment_value",
                        "weight_col", "cluster_col", "panel_id_col", "strata"):
                val = design.get(key)
                if val in (None, "", [], {}):
                    continue
                r = d_table.add_row().cells
                r[0].text = key.replace("_", " ").title()
                r[1].text = ", ".join(map(str, val)) if isinstance(val, list) else str(val)
            pairs = design.get("pre_post_pairs") or []
            if pairs:
                doc.add_paragraph("Pre/Post pairs:").runs[0].bold = True
                for p in pairs:
                    if isinstance(p, dict):
                        doc.add_paragraph(f"  • {p.get('pre','?')} → {p.get('post','?')}",
                                          style="List Bullet")

        # ── Variable directory ──
        doc.add_heading("Variables", level=1)
        for idx, col in enumerate(df.columns, start=1):
            role_meta = roles.get(col, {}) or {}
            col_type = _detect_type(df[col], overrides.get(col))
            scale = role_meta.get("scale", "—")
            role = role_meta.get("role", "—")
            paired_with = role_meta.get("paired_with")
            mr_set = role_meta.get("mr_set_id")
            units = role_meta.get("units")
            labels = role_meta.get("value_labels") or {}

            # Variable header
            ph = doc.add_paragraph()
            run = ph.add_run(f"{idx}. {col}")
            run.bold = True
            run.font.size = Pt(12)

            # Meta table
            t = doc.add_table(rows=0, cols=2)
            t.style = "Light List Accent 1"
            base_rows = [
                ("Type", col_type),
                ("Scale", scale),
                ("Role", role),
                ("N (non-missing)", f"{len(df) - _missing_count(df[col]):,}"),
                ("Missing", f"{_missing_count(df[col]):,} ({_missing_count(df[col])/max(1,len(df))*100:.1f}%)"),
                ("Unique values", f"{df[col].nunique(dropna=True):,}"),
            ]
            if units:
                base_rows.append(("Units", str(units)))
            if paired_with:
                base_rows.append(("Paired with", str(paired_with)))
            if mr_set:
                base_rows.append(("MR set", str(mr_set)))
            if role_meta.get("benchmark_link"):
                base_rows.append(("Benchmark", str(role_meta["benchmark_link"])))

            for k, v in base_rows:
                r = t.add_row().cells
                r[0].text = k
                r[1].text = v
                for cell in r:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(10)

            # Numeric summary
            if col_type == "numeric":
                ns = _numeric_summary(df[col])
                if ns.get("n", 0) > 0:
                    p = doc.add_paragraph()
                    p.add_run("Descriptives: ").bold = True
                    p.add_run(
                        f"mean={ns['mean']:.3f}, SD={ns['sd']:.3f}, "
                        f"min={ns['min']:.3f}, max={ns['max']:.3f}, "
                        f"median={ns['median']:.3f} (Q1={ns['q1']:.3f}, Q3={ns['q3']:.3f})"
                    )

            # Value labels block
            if labels:
                doc.add_paragraph("Value labels:").runs[0].bold = True
                lt = doc.add_table(rows=1, cols=2)
                lt.style = "Light Grid Accent 1"
                hdr = lt.rows[0].cells
                hdr[0].text = "Code"
                hdr[1].text = "Label"
                for code, lbl in labels.items():
                    rr = lt.add_row().cells
                    rr[0].text = str(code)
                    rr[1].text = str(lbl)

            # Frequencies for text / low-cardinality numeric
            if config.include_frequencies:
                n_unique = df[col].nunique(dropna=True)
                show_freq = (col_type == "text" and n_unique <= config.max_freq_levels * 4) or \
                            (col_type == "numeric" and n_unique <= config.max_freq_levels)
                if show_freq and n_unique > 0:
                    doc.add_paragraph("Frequencies:").runs[0].bold = True
                    ft = doc.add_table(rows=1, cols=3)
                    ft.style = "Light Grid Accent 1"
                    hdr = ft.rows[0].cells
                    hdr[0].text = "Value"
                    hdr[1].text = "Count"
                    hdr[2].text = "Percent"
                    for val, cnt, pct in _frequency(df[col], config.max_freq_levels):
                        rr = ft.add_row().cells
                        rr[0].text = val
                        rr[1].text = f"{cnt:,}"
                        rr[2].text = f"{pct:.1f}%"
                        for cell in rr:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.font.size = Pt(9)

            doc.add_paragraph()  # spacer

        # ── Write ──
        fname = config.filename or f"codebook_{config.dataset_id[:8]}.docx"
        if not fname.lower().endswith(".docx"):
            fname += ".docx"
        out_path = EXPORTS_DIR / fname
        doc.save(str(out_path))

        add_audit_log(config.dataset_id, "export_codebook",
                      f"Codebook with {len(df.columns)} variables exported to {fname}")

        return {
            "path": str(out_path),
            "message": f"Codebook exported to {fname}",
            "download_filename": fname,
            "n_variables": len(df.columns),
        }
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(500, f"Codebook export failed: {e}")
