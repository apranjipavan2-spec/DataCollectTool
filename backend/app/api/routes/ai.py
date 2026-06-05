from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from fastapi.responses import Response
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.core.deps import require_role
from app.models.system_setting import SystemSetting
from app.models.form import Form
from app.models.submission import Submission
from app.services import ai_service

router = APIRouter()
require_org_admin   = require_role("org_admin")
require_supervisor  = require_role("org_admin", "supervisor")
require_master      = require_role("master_admin")


def _get_global_ai_cfg(db: Session) -> dict:
    """Return cfg in legacy format {provider, api_key, model} for ai_service compatibility."""
    row = db.query(SystemSetting).filter(SystemSetting.key == "ai_config").first()
    if not row:
        return {}
    cfg = row.value or {}
    # New multi-key format → resolve active provider
    if "keys" in cfg:
        active = cfg.get("active_provider", "")
        key_cfg = cfg.get("keys", {}).get(active, {})
        return {"provider": active, "api_key": key_cfg.get("api_key", ""), "model": key_cfg.get("model", "")}
    # Legacy single-key format
    return cfg


def _logged_cfg(db: Session, user: dict, feature: str) -> dict:
    """Get global AI cfg with usage-log context attached for this user + feature."""
    return ai_service.with_log(
        _get_global_ai_cfg(db),
        tenant_id=user.get("tenant_id"),
        user_id=user.get("sub"),
        feature=feature,
    )


# ── AI Config (global — set by master_admin, visible to all) ──────────────

PROVIDER_DEFAULTS = {"openai": "gpt-4o", "anthropic": "claude-sonnet-4-6", "gemini": "gemini-2.0-flash", "deepseek": "deepseek-v4-flash"}

@router.get("/config")
def get_ai_config(user: dict = Depends(require_org_admin), db: Session = Depends(get_db)):
    row = db.query(SystemSetting).filter(SystemSetting.key == "ai_config").first()
    cfg = row.value if row else {}

    ALL_PROVIDERS = ["openai", "anthropic", "gemini", "deepseek"]
    if "keys" in cfg:
        keys_status = {
            p: {"configured": bool(cfg["keys"].get(p, {}).get("api_key")),
                "model": cfg["keys"].get(p, {}).get("model", PROVIDER_DEFAULTS.get(p, ""))}
            for p in ALL_PROVIDERS
        }
        active = cfg.get("active_provider", "")
        return {
            "active_provider": active,
            "keys": keys_status,
            "configured": bool(cfg["keys"].get(active, {}).get("api_key")),
            "provider": active,
        }

    # Legacy format
    return {
        "active_provider": cfg.get("provider", ""),
        "provider": cfg.get("provider", ""),
        "keys": {p: {"configured": cfg.get("provider") == p and bool(cfg.get("api_key")),
                     "model": PROVIDER_DEFAULTS.get(p, "")} for p in ALL_PROVIDERS},
        "configured": bool(cfg.get("provider") and cfg.get("api_key")),
    }


@router.patch("/config")
def update_ai_config(body: dict, user: dict = Depends(require_master), db: Session = Depends(get_db)):
    """master_admin sets per-provider keys and chooses active provider."""
    row = db.query(SystemSetting).filter(SystemSetting.key == "ai_config").first()
    if not row:
        row = SystemSetting(key="ai_config", value={"keys": {}, "active_provider": ""})
        db.add(row)

    cfg = dict(row.value or {})
    if "keys" not in cfg:
        # Migrate legacy format
        cfg = {"keys": {}, "active_provider": cfg.get("provider", "")}

    provider = body.get("provider", "")
    if provider and body.get("api_key"):
        cfg["keys"].setdefault(provider, {})
        cfg["keys"][provider]["api_key"] = body["api_key"]
        cfg["keys"][provider]["model"] = body.get("model") or PROVIDER_DEFAULTS.get(provider, "")
    if body.get("active_provider"):
        cfg["active_provider"] = body["active_provider"]
    elif provider:
        cfg["active_provider"] = provider

    from sqlalchemy.orm.attributes import flag_modified
    row.value = cfg
    flag_modified(row, "value")
    db.commit()
    return {"ok": True, "active_provider": cfg.get("active_provider")}


# ── AI Usage Tracking ─────────────────────────────────────────────────────

@router.get("/usage")
def get_ai_usage(
    days: int = 30,
    feature: str = "",
    user: dict = Depends(require_org_admin),
    db: Session = Depends(get_db),
):
    """Per-user AI usage totals. master_admin sees all tenants; org_admin sees own tenant only.
    Query params: days (default 30, max 365), feature (optional filter)."""
    from sqlalchemy import func as sa_func, case as sa_case
    from app.models.ai_usage_log import AiUsageLog
    from app.models.user import User
    from datetime import datetime, timedelta, timezone

    days = max(1, min(int(days or 30), 365))
    since = datetime.now(timezone.utc) - timedelta(days=days)

    q = db.query(
        AiUsageLog.user_id,
        AiUsageLog.tenant_id,
        sa_func.count(AiUsageLog.id).label("calls"),
        sa_func.sum(AiUsageLog.tokens_in).label("tokens_in"),
        sa_func.sum(AiUsageLog.tokens_out).label("tokens_out"),
        sa_func.sum(sa_case((AiUsageLog.success == False, 1), else_=0)).label("errors"),
        sa_func.max(AiUsageLog.created_at).label("last_call"),
    ).filter(AiUsageLog.created_at >= since)

    if user.get("role") != "master_admin":
        q = q.filter(AiUsageLog.tenant_id == user["tenant_id"])
    if feature:
        q = q.filter(AiUsageLog.feature == feature)

    rows = q.group_by(AiUsageLog.user_id, AiUsageLog.tenant_id).all()

    # Hydrate user/tenant names
    user_ids = [r.user_id for r in rows if r.user_id]
    users = {str(u.id): u for u in db.query(User).filter(User.id.in_(user_ids)).all()} if user_ids else {}

    result = []
    for r in rows:
        u = users.get(str(r.user_id)) if r.user_id else None
        result.append({
            "user_id": str(r.user_id) if r.user_id else None,
            "user_name": (u.name if u else None) or (u.email if u else None) or "(unknown)",
            "user_email": u.email if u else None,
            "tenant_id": str(r.tenant_id),
            "calls": int(r.calls or 0),
            "tokens_in": int(r.tokens_in or 0),
            "tokens_out": int(r.tokens_out or 0),
            "errors": int(r.errors or 0),
            "last_call": r.last_call.isoformat() if r.last_call else None,
        })
    result.sort(key=lambda x: x["calls"], reverse=True)

    # Per-feature breakdown for the same window/scope
    fq = db.query(
        AiUsageLog.feature,
        sa_func.count(AiUsageLog.id).label("calls"),
        sa_func.sum(AiUsageLog.tokens_in).label("tokens_in"),
        sa_func.sum(AiUsageLog.tokens_out).label("tokens_out"),
    ).filter(AiUsageLog.created_at >= since)
    if user.get("role") != "master_admin":
        fq = fq.filter(AiUsageLog.tenant_id == user["tenant_id"])
    by_feature = [
        {
            "feature": fr.feature,
            "calls": int(fr.calls or 0),
            "tokens_in": int(fr.tokens_in or 0),
            "tokens_out": int(fr.tokens_out or 0),
        }
        for fr in fq.group_by(AiUsageLog.feature).order_by(sa_func.count(AiUsageLog.id).desc()).all()
    ]

    return {
        "days": days,
        "scope": "global" if user.get("role") == "master_admin" else "tenant",
        "per_user": result,
        "by_feature": by_feature,
    }


# ── AI Features ───────────────────────────────────────────────────────────

@router.post("/report/{form_id}")
async def generate_report(
    form_id: str,
    background_tasks: BackgroundTasks,
    user: dict = Depends(require_supervisor),
    db: Session = Depends(get_db),
):
    """Start async AI report generation. Returns {job_id}. Poll GET /fg/ai-jobs/{job_id}."""
    from app.models.user_tool_project import UserToolProject
    from sqlalchemy.orm.attributes import flag_modified

    if user.get("role") != "master_admin":
        from app.api.routes.billing import check_feature
        check_feature(user["tenant_id"], "ai_writer", db)

    cfg = _logged_cfg(db, user, "report")
    if not cfg.get("api_key"):
        raise HTTPException(400, "AI not configured. Contact your platform administrator.")

    form = db.query(Form).filter(Form.id == form_id, Form.tenant_id == user["tenant_id"]).first()
    if not form:
        raise HTTPException(404, "Form not found")

    subs = db.query(Submission).filter(
        Submission.form_id == form_id, Submission.tenant_id == user["tenant_id"]
    ).limit(50).all()
    field_labels = [
        f.get("label", f.get("id", ""))
        for section in (form.json_schema or {}).get("sections", [])
        for f in section.get("fields", [])
    ]
    sub_data = [s.data_json for s in subs if s.data_json]
    form_title = form.title or "Survey"

    # Create job record
    import uuid as _uuid_mod
    job = UserToolProject(
        tenant_id=user["tenant_id"], user_id=user["sub"],
        tool="ai_job", name=f"Report: {form_title}",
        data={"status": "pending", "steps": [], "result": None, "error": None},
    )
    db.add(job); db.commit(); db.refresh(job)
    job_id = str(job.id)

    def _upd(sess, status, step=None, result=None, error=None):
        rec = sess.query(UserToolProject).filter(UserToolProject.id == job_id).first()
        if not rec: return
        d = dict(rec.data or {})
        d["status"] = status
        if step: d.setdefault("steps", []).append(step)
        if result is not None: d["result"] = result
        if error  is not None: d["error"]  = error
        rec.data = d; flag_modified(rec, "data"); sess.commit()

    job_tenant_id = str(user["tenant_id"])

    async def _run():
        from app.core.database import SessionLocal, set_tenant_context
        with SessionLocal() as sess:
            set_tenant_context(sess, job_tenant_id)
            _upd(sess, "running", f"Loaded {len(sub_data)} submissions — calling AI model…")
            try:
                report_md = await ai_service.generate_report(cfg, form_title, field_labels, sub_data)
                _upd(sess, "done", "Report complete!", result=report_md)
            except Exception as e:
                _upd(sess, "failed", error=str(e))

    background_tasks.add_task(_run)
    return {"job_id": job_id}


@router.post("/suggest-skip-logic")
async def suggest_skip_logic(body: dict, user: dict = Depends(require_org_admin), db: Session = Depends(get_db)):
    if user.get("role") != "master_admin":
        from app.api.routes.billing import check_feature
        check_feature(user["tenant_id"], "ai_cleaning", db)
    cfg = _logged_cfg(db, user, "skip_logic")
    try:
        suggestions = await ai_service.suggest_skip_logic(cfg, body.get("question_text", ""), body.get("form_fields", []), body.get("user_description", ""))
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(503, f"AI provider error: {e}")
    return {"suggestions": suggestions}


@router.post("/translate")
async def translate_labels(body: dict, user: dict = Depends(require_org_admin), db: Session = Depends(get_db)):
    if user.get("role") != "master_admin":
        from app.api.routes.billing import check_feature
        check_feature(user["tenant_id"], "ai_cleaning", db)
    cfg = _logged_cfg(db, user, "translate")
    try:
        translated = await ai_service.translate_labels(cfg, body.get("labels", []), body.get("target_lang", "Hindi"))
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(503, f"AI provider error: {e}")
    return {"translated": translated}


@router.post("/writer")
async def ai_writer(body: dict, user: dict = Depends(require_supervisor), db: Session = Depends(get_db)):
    if user.get("role") != "master_admin":
        from app.api.routes.billing import check_feature
        check_feature(user["tenant_id"], "ai_writer", db)
    cfg = _logged_cfg(db, user, "writer")
    try:
        report_md = await ai_service.generate_styled_report(
            cfg=cfg,
            style=body.get("style", "field_survey"),
            form_title=body.get("form_title", "Survey"),
            date_range=body.get("date_range", ""),
            sample_size=body.get("sample_size", 0),
            table_data=body.get("table_data", ""),
            chart_descriptions=body.get("chart_descriptions", ""),
            custom_context=body.get("custom_context", ""),
        )
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(503, f"AI error: {e}")
    return {"report_md": report_md}


@router.post("/writer/export-docx")
def export_docx(body: dict, user: dict = Depends(require_supervisor)):
    import io, re
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(503, "python-docx not installed")

    md = body.get("report_md", "")
    title = body.get("title", "Report")
    doc = Document()
    t = doc.add_heading(title, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in md.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph('')
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        else:
            p = doc.add_paragraph()
            for part in re.split(r'(\*\*.*?\*\*)', line):
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{title}.docx"'},
    )


# ── AI Form Generation ────────────────────────────────────────────────────

@router.post("/generate-form")
async def generate_form(
    body: dict,
    background_tasks: BackgroundTasks,
    user: dict = Depends(require_org_admin),
    db: Session = Depends(get_db),
):
    """
    Generate a complete form schema from study objectives.
    Returns immediately with form_id + status='pending'.
    Client polls GET /forms/{form_id}/generation-status.
    Body: { objectives, study_type, reference_form_ids?, notify_email? }
    """
    import uuid as _uuid
    from app.core.database import SessionLocal

    cfg = _logged_cfg(db, user, "generate_form")
    if not cfg.get("api_key"):
        raise HTTPException(400, "AI not configured. Contact your platform administrator.")

    objectives = body.get("objectives", "").strip()
    if not objectives:
        raise HTTPException(400, "objectives is required")

    study_type = body.get("study_type", "field_survey")
    reference_form_ids = body.get("reference_form_ids", [])

    # Fetch reference form schemas for context (max 3 to keep prompt manageable)
    ref_schemas = []
    for fid in reference_form_ids[:3]:
        ref = db.query(Form).filter(Form.id == fid, Form.tenant_id == user["tenant_id"]).first()
        if ref and ref.json_schema:
            ref_schemas.append({"title": ref.title, "schema": ref.json_schema})

    # Create placeholder form with pending status
    new_form = Form(
        tenant_id=user["tenant_id"],
        title=f"[Generating…] {study_type.replace('_', ' ').title()} Form",
        json_schema={"title": "", "sections": [], "version": 1},
        version=1,
        status="draft",
        created_by=user.get("sub"),
        generation_status="pending",
    )
    db.add(new_form)
    db.commit()
    db.refresh(new_form)
    form_id = str(new_form.id)
    tenant_id = str(user["tenant_id"])

    async def _run():
        from app.core.database import SessionLocal, set_tenant_context
        with SessionLocal() as sess:
            set_tenant_context(sess, tenant_id)
            form_rec = sess.query(Form).filter(Form.id == form_id, Form.tenant_id == tenant_id).first()
            if not form_rec:
                return
            try:
                result = await ai_service.generate_form_schema(
                    cfg=cfg,
                    objectives=objectives,
                    study_type=study_type,
                    reference_schemas=ref_schemas,
                )
                form_rec.json_schema = result["schema"]
                form_rec.title = result.get("title", "Generated Form")
                form_rec.generation_status = "done"
                form_rec.generation_error = None
            except Exception as e:
                form_rec.generation_status = "failed"
                form_rec.generation_error = str(e)
            sess.commit()

    background_tasks.add_task(_run)
    return {"form_id": form_id, "status": "pending", "message": "Form generation started. Check back in ~1–2 minutes."}
