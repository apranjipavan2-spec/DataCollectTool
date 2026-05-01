"""Scheduled report delivery service.

Generates AI-written reports from saved tabulations and emails them as DOCX
attachments to configured recipients.
"""
import io
import logging
import re
from datetime import datetime, timezone

from sqlalchemy.orm import Session

from app.services.email import send_email

logger = logging.getLogger(__name__)


def _md_to_docx_bytes(md: str, title: str) -> bytes:
    """Convert markdown to DOCX bytes using python-docx."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    t = doc.add_heading(title, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in md.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        else:
            p = doc.add_paragraph()
            for part in re.split(r"(\*\*.*?\*\*)", line):
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _should_run_now(sr, now: datetime) -> bool:
    """Return True if this scheduled report is due to run at `now`."""
    hour = int(sr.send_hour or 7)
    if now.hour != hour:
        return False

    freq = sr.frequency or "weekly"
    if freq == "daily":
        return True
    if freq == "weekly":
        dow = int(sr.send_dow) if sr.send_dow is not None else 0
        return now.weekday() == dow
    if freq == "monthly":
        return now.day == 1
    return False


def run_scheduled_reports(db: Session) -> dict:
    """Check all active scheduled reports and deliver any that are due now."""
    import asyncio
    from app.models.scheduled_report import ScheduledReport
    from app.models.program import Program
    from app.models.program_analysis import ProgramAnalysis
    from app.services import ai_service
    from app.core.database import set_tenant_context

    now = datetime.now(timezone.utc)
    sent = 0
    skipped = 0
    errors = 0

    reports = db.query(ScheduledReport).filter(ScheduledReport.is_active == True).all()
    for sr in reports:
        if not _should_run_now(sr, now):
            skipped += 1
            continue
        # Skip if already sent within the past hour (prevents duplicate delivery)
        if sr.last_sent_at and (now - sr.last_sent_at).total_seconds() < 3600:
            skipped += 1
            continue

        try:
            set_tenant_context(db, str(sr.tenant_id))

            # Load program name
            prog = db.query(Program).filter(Program.id == sr.program_id).first() if sr.program_id else None
            prog_name = prog.name if prog else "Report"

            # Load saved tabulations from most recent manual analysis
            table_data = ""
            if sr.program_id:
                analysis = (
                    db.query(ProgramAnalysis)
                    .filter(
                        ProgramAnalysis.program_id == sr.program_id,
                        ProgramAnalysis.tenant_id == sr.tenant_id,
                        ProgramAnalysis.source == "manual",
                    )
                    .order_by(ProgramAnalysis.created_at.desc())
                    .first()
                )
                if analysis and analysis.table_configs:
                    rows_summary = []
                    for tab in (analysis.table_configs or [])[:10]:
                        rows_summary.append(f"**{tab.get('title','Table')}** ({tab.get('aggregation','count')}): " +
                            ", ".join(f"{r.get('group','?')}: {r.get('value','?')}" for r in (tab.get('rows') or [])[:5]))
                    table_data = "\n".join(rows_summary)

            if not table_data:
                logger.info("[ScheduledReports] No tabulations for report %s — skipping", sr.id)
                skipped += 1
                continue

            # Get AI config
            from app.models.ai_config import AiConfig
            cfg_row = db.query(AiConfig).filter(AiConfig.tenant_id == None).first()
            cfg = cfg_row.config if cfg_row else {}

            # Generate report markdown
            report_md = asyncio.run(ai_service.generate_styled_report(
                cfg=cfg,
                style=sr.style or "field_survey",
                form_title=prog_name,
                date_range=now.strftime("%B %Y"),
                sample_size=0,
                table_data=table_data,
                chart_descriptions="",
                custom_context=sr.custom_context or "",
            ))

            # Convert to DOCX
            title = f"{prog_name} — {now.strftime('%B %Y')}"
            docx_bytes = _md_to_docx_bytes(report_md, title)
            filename = f"{prog_name.replace(' ', '_')}_{now.strftime('%Y_%m')}.docx"

            html_body = f"""
            <div style="font-family: system-ui, sans-serif; max-width: 600px; margin: 0 auto; padding: 24px;">
              <h2 style="color: #1e1e2e;">{sr.name}</h2>
              <p>Your scheduled report for <strong>{prog_name}</strong> is attached.</p>
              <p style="color: #888; font-size: 13px;">Generated: {now.strftime('%d %b %Y, %H:%M UTC')}</p>
              <hr style="border: none; border-top: 1px solid #eee; margin: 20px 0;">
              <p style="color: #aaa; font-size: 12px;">This is an automated report from FieldGovern.
              To change your schedule, log in and visit FG Writer → Schedule.</p>
            </div>
            """

            recipients = sr.recipient_emails or []
            for email in recipients:
                ok = send_email(
                    to=email,
                    subject=f"[FieldGovern] {sr.name} — {now.strftime('%B %Y')}",
                    html_body=html_body,
                    attachment_bytes=docx_bytes,
                    attachment_filename=filename,
                )
                if ok:
                    sent += 1

            sr.last_sent_at = now
            db.commit()
            logger.info("[ScheduledReports] Delivered '%s' to %d recipients", sr.name, len(recipients))

        except Exception:
            logger.exception("[ScheduledReports] Error delivering report %s", sr.id)
            errors += 1

    return {"sent": sent, "skipped": skipped, "errors": errors}
